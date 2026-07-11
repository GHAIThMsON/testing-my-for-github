from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

title = doc.add_heading('Baity Android App - Security Findings Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Date: {datetime.date.today()}  |  Package: com.baity.android  |  Version: 3.3.7')
doc.add_paragraph('')

doc.add_heading('Scope', level=2)
doc.add_paragraph('Filtered findings: only real exploitable vulnerabilities. Theoretical/low-likelihood items removed.')
doc.add_paragraph('')

# ========================================================================
# SECTION 1: DIRECT IMPACT
# ========================================================================
doc.add_heading('SECTION 1: DIRECT IMPACT - Exploitable Now', level=1)

# 1
doc.add_heading('1. FileProvider Exposes Entire Filesystem', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Critical')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: The FileProvider (com.crazecoder.openfile.FileProvider) exposes root-path, files-path, external-path, external-cache-path, and external-files-path all with path="." - meaning the entire filesystem is accessible via content:// URIs.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Any app on the device (or ADB shell) can read/write any file belonging to Baity. This includes SQLite databases with user data, SharedPreferences with tokens, and cached files. Nearly all local storage is readable by any third-party app.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('adb shell content query --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/ --projection _display_name:_size', style='List Bullet')
doc.add_paragraph('adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/shared_prefs/FlutterSharedPreferences.xml', style='List Bullet')
doc.add_paragraph('adb shell content read --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/root/data/data/com.baity.android/databases/OneSignal.db', style='List Bullet')
doc.add_paragraph('adb shell content query --uri content://com.baity.android.fileProvider.com.crazecoder.openfile/files_path/ --projection _display_name:_size', style='List Bullet')
doc.add_paragraph('')

# 2
doc.add_heading('2. Cleartext HTTP Traffic Fully Permitted', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Critical')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: Both android:usesCleartextTraffic="true" in the manifest AND the network_security_config.xml set cleartextTrafficPermitted=true with zero domain restrictions. ALL HTTP traffic is allowed.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: An attacker on the same network (Wi-Fi, cellular, VPN) can perform ARP spoofing or DNS poisoning to intercept all HTTP traffic. Since no SSL pinning exists either, even HTTPS can be bypassed with a proxy. User credentials, auth tokens, and API responses can all be captured or modified in transit.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Configure Burp Suite as proxy on emulator (Settings > Wi-Fi > Proxy > Manual)', style='List Bullet')
doc.add_paragraph('Install Burp CA certificate on the emulator', style='List Bullet')
doc.add_paragraph('Launch the app and perform login, search, reservation flows', style='List Bullet')
doc.add_paragraph('Observe all HTTP/HTTPS calls in Burp - confirm cleartext and no pinning', style='List Bullet')
doc.add_paragraph('')

# 3
doc.add_heading('3. No SSL Pinning Implemented', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: No certificate or public key pinning for any API endpoint (v3.ibaity.com, stagingapi.ibaity.com, onesignal.com, sentry.io).')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Combined with cleartext being permitted, an attacker with network access can intercept all HTTPS traffic using a proxy CA certificate. The app trusts any valid certificate. This enables credential theft, session hijacking, and data manipulation on any network.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Set Burp Suite as device proxy', style='List Bullet')
doc.add_paragraph('Burp CA cert will be accepted by the system trust store (no pinning bypass needed)', style='List Bullet')
doc.add_paragraph('Intercept all API calls to v3.ibaity.com/api and verify full request/response visibility', style='List Bullet')
doc.add_paragraph('')

# 4
doc.add_heading('4. Hardcoded Google API Keys (Potentially Unrestricted)', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: Google Maps API key AIzaSyCVOlkpjQzMEiUuz3Q0tavNv7kJOxISP-s is in AndroidManifest.xml. Google API key AIzaSyAku5tBaHh4UxZ2FbwieVfM95Jy9NR4mqo is in strings.xml. Both in plaintext.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: If unrestricted (no package name/SHA-1 binding), anyone can extract and use these keys. The Maps key could serve maps on unauthorized sites at $7 per 1000 loads. The generic API key could access Google Cloud services, potentially incurring significant costs or accessing project resources if misconfigured.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Extract keys: grep -r "AIza" apktool_out/', style='List Bullet')
doc.add_paragraph('Test Maps key: curl "https://maps.googleapis.com/maps/api/geocode/json?address=Baghdad&key=AIzaSyCVOlkpjQzMEiUuz3Q0tavNv7kJOxISP-s"', style='List Bullet')
doc.add_paragraph('Test generic key: curl "https://www.googleapis.com/oauth2/v1/tokeninfo?access_token=AIzaSyAku5tBaHh4UxZ2FbwieVfM95Jy9NR4mqo"', style='List Bullet')
doc.add_paragraph('If either returns valid data, the key is unrestricted', style='List Bullet')
doc.add_paragraph('')

# 5
doc.add_heading('5. Sentry DSN Exposed in Binary', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: Sentry DSN https://84a2aabbf750391b0127349e161f1322@o4510570135748608.ingest.de.sentry.io/4510570142695504 is hardcoded in the compiled Dart code (libapp.so).')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Anyone with this DSN can send fake error events, flood the project (DoS), or inspect error data. If Sentry captures request bodies or user data in breadcrumbs, this exposes what user information leaks through errors. The DSN public key (first hash segment) is a credential that authenticates to the Sentry project.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Extract DSN: strings libapp.so | grep sentry.io', style='List Bullet')
doc.add_paragraph('Send fake event: curl -X POST "https://o4510570135748608.ingest.de.sentry.io/api/4510570142695504/store/" -d \'{"message":"test_impact","level":"info"}\'', style='List Bullet')
doc.add_paragraph('If accepted (201 response), the DSN is active and writable by anyone', style='List Bullet')
doc.add_paragraph('')

# 6
doc.add_heading('6. Exported SignInWithAppleCallback Activity', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: High')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: The activity com.aboutyou.dart_packages.sign_in_with_apple.SignInWithAppleCallback is exported and handles signinwithapple://callback deep links.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Any app on the device can send intents to this activity. A malicious app can craft intents carrying fake Apple Sign-In response data. If the app trusts the callback data without verifying the source, an attacker could hijack the Apple Sign-In flow and impersonate users. This is a classic OAuth callback interception vulnerability.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('adb shell am start -a android.intent.action.VIEW -d "signinwithapple://callback?code=test123&id_token=fake.jwt.token&state=abc" com.baity.android', style='List Bullet')
doc.add_paragraph('Check if the activity logs the callback or processes it', style='List Bullet')
doc.add_paragraph('Build a PoC app that registers the same intent filter and intercepts real callbacks', style='List Bullet')
doc.add_paragraph('Monitor logcat: adb logcat | grep -i apple', style='List Bullet')
doc.add_paragraph('')

# 7
doc.add_heading('7. Facebook Codeless Debug Logging Enabled in Production', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: The meta-data com.facebook.sdk.CodelessDebugLogEnabled is set to true in the production manifest.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Facebook SDK outputs detailed UI event data, user interaction details, and SDK internals to logcat. Any app with READ_LOGS or via ADB can read this. This includes which buttons users tap, what screens they view, and potentially sensitive interaction data. In a real estate app, this could reveal property interests, saved searches, and communication patterns.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('adb logcat | Select-String -Pattern "(facebook|FBAnalytics|Codeless)"', style='List Bullet')
doc.add_paragraph('Perform actions in the app (login, search, view listings) and observe Facebook logs', style='List Bullet')
doc.add_paragraph('')

# 8
doc.add_heading('8. Google Tag Manager Preview Mode Exported', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: TagManagerPreviewActivity is exported with scheme tagmanager.c.com.baity.android.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Any app can trigger GTM preview mode, potentially injecting modified container values. This could alter analytics, remote config, Firebase feature flags, and A/B test assignments. If GTM controls backend URLs or feature toggles, this could redirect API calls or enable hidden functionality.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('adb shell am start -a android.intent.action.VIEW -d "tagmanager.c.com.baity.android" com.baity.android', style='List Bullet')
doc.add_paragraph('Check if a preview/overlay appears indicating GTM preview mode', style='List Bullet')
doc.add_paragraph('')

# ========================================================================
# SECTION 2: ABUSE POTENTIAL
# ========================================================================
doc.add_heading('SECTION 2: ABUSE POTENTIAL - May Lead to Impact Under Conditions', level=1)

# 9
doc.add_heading('9. UXCam Session Recording Captures User Screens', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: UXCam SDK (v3.8.3) with multi-session recording, auto screen tagging, crash handling, and honorFlagSecure=false.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Records full user sessions including screenshots. If a user enters personal data, views documents, or performs financial transactions, these are recorded and uploaded to UXCam servers. An attacker who accesses the UXCam dashboard (via credential leak, SSRF, or phishing) gains full session replays. The honorFlagSecure=false setting means even screens protected with FLAG_SECURE can be recorded.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Intercept network traffic with Burp, filter for uxcam.com domains', style='List Bullet')
doc.add_paragraph('Check for session uploads with image/screenshot data', style='List Bullet')
doc.add_paragraph('adb logcat | grep -i uxcam', style='List Bullet')
doc.add_paragraph('')

# 10
doc.add_heading('10. Task Hijacking via OneSignal Notification Activities', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: OneSignal\'s NotificationOpenedActivity has android:taskAffinity="" and is exported. This enables task hijacking attacks.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: A malicious app with the same taskAffinity can inject its own activity into the Baity app task stack. When a user taps a notification, the attacker\'s activity appears instead of the legitimate one. This enables phishing: showing a fake login screen within the Baity app context. Users see the Baity icon in the task switcher but interact with the attacker\'s UI.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Create a PoC app with taskAffinity="" and exported activity', style='List Bullet')
doc.add_paragraph('Install it, then trigger a OneSignal notification', style='List Bullet')
doc.add_paragraph('Observe if the malicious activity appears instead of the legitimate one', style='List Bullet')
doc.add_paragraph('adb shell dumpsys activity activities | grep -i baity', style='List Bullet')
doc.add_paragraph('')

# 11
doc.add_heading('11. TikTok SDK Collecting User Data', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: TikTok Events SDK with anonymous ID tracking. Data flows to analytics.us.tiktok.com. Email/phone regex patterns present in SDK config.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: User behavior data is sent to ByteDance (TikTok) servers. The SDK includes pattern matching for emails and phone numbers - could be matching users across apps for ad attribution. Privacy implications: users may not be aware TikTok is tracking their real estate browsing. Data includes device model, OS, app version, and user interactions.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Intercept traffic: Burp filter for tiktok.com domains', style='List Bullet')
doc.add_paragraph('Check stored data: cat /data/data/com.baity.android/shared_prefs/com.tiktok.sdk.keystore.xml', style='List Bullet')
doc.add_paragraph('adb logcat | grep -i tiktok', style='List Bullet')
doc.add_paragraph('')

# 12
doc.add_heading('12. SYSTEM_ALERT_WINDOW Enables Overlay Attacks', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Medium')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: The app requests SYSTEM_ALERT_WINDOW permission for drawing overlays on top of other apps.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Once granted, a malicious component within the app (or another app with same permission) can draw transparent overlays over UI elements. Combined with exported activities, this enables clickjacking: an invisible overlay over "Delete Listing" could be triggered when the user thinks they are tapping "Save". The user\'s action is hijacked without their knowledge.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Check grant status: adb shell appops get com.baity.android SYSTEM_ALERT_WINDOW', style='List Bullet')
doc.add_paragraph('Create a proof-of-concept overlay app that draws an invisible button over a critical action in Baity', style='List Bullet')
doc.add_paragraph('')

# 13
doc.add_heading('13. FCM Push Tokens Stored Unencrypted', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Low-Moderate')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: Firebase Cloud Messaging tokens stored in plaintext in com.google.android.gms.appid.xml [line 3-4].')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: FCM tokens are device identifiers for push. If extracted (via FileProvider vuln above), an attacker can send spoofed push notifications to the device. Combined with OneSignal notification channels (Real Estates, Account categories exist), the attacker could craft realistic-looking notifications to phish credentials or trick users into taking actions.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Read stored tokens via FileProvider vuln or: adb shell cat /data/data/com.baity.android/shared_prefs/com.google.android.gms.appid.xml', style='List Bullet')
doc.add_paragraph('Use extracted token with Firebase Cloud Messaging API to send a test push', style='List Bullet')
doc.add_paragraph('')

# 14
doc.add_heading('14. HTTP Scheme in Auto-Verified Deep Links', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Low-Moderate')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
doc.add_paragraph('')
doc.add_paragraph('What: MainActivity intent-filter with autoVerify=true includes both https:// and http:// schemes for ibaity.com.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: On a non-HTTPS network, an attacker can intercept HTTP traffic to ibaity.com and inject a redirect or intent with crafted parameters. If the app processes deep link parameters (like tokens, action IDs) from the URL without validation, this could enable XSS via WebView, account linking attacks, or unauthorized state changes.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('adb shell am start -a android.intent.action.VIEW -d "http://ibaity.com/?action=delete_listing&id=123" com.baity.android', style='List Bullet')
doc.add_paragraph('adb shell am start -a android.intent.action.VIEW -d "http://ibaity.com/?token=ATTACKER_TOKEN" com.baity.android', style='List Bullet')
doc.add_paragraph('Observe if the app processes these parameters', style='List Bullet')
doc.add_paragraph('')

# 15
doc.add_heading('15. ClickUp Project Management Endpoints in Production', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: Low')
run.bold = True
doc.add_paragraph('')
doc.add_paragraph('What: ClickUp API endpoints https://api.clickup.com/api/v2/list/901814381791/task and /api/v2/task/ are present in the production binary.')
doc.add_paragraph('')
doc.add_paragraph('Why it matters: Internal project management URLs should not be in production apps. Suggests development features compiled into release. While ClickUp requires authentication, the list ID 901814381791 is exposed. If a token is ever stored insecurely, this endPoint allows direct task manipulation.')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('How to test:')
run.bold = True
doc.add_paragraph('Monitor network: check if app calls api.clickup.com during normal use', style='List Bullet')
doc.add_paragraph('adb logcat | grep -i clickup', style='List Bullet')
doc.add_paragraph('')


# ========================================================================
# SUMMARY TABLE
# ========================================================================
doc.add_heading('Summary', level=1)
table = doc.add_table(rows=16, cols=4)
table.style = 'Light Grid Accent 1'

cells = table.rows[0].cells
cells[0].text = '#'
cells[1].text = 'Finding'
cells[2].text = 'Severity'
cells[3].text = 'Category'

findings = [
    ('1', 'FileProvider Full Filesystem Access', 'CRITICAL', 'Direct Impact'),
    ('2', 'Cleartext HTTP Traffic Permitted', 'CRITICAL', 'Direct Impact'),
    ('3', 'No SSL Pinning', 'HIGH', 'Direct Impact'),
    ('4', 'Hardcoded Google API Keys', 'HIGH', 'Direct Impact'),
    ('5', 'Sentry DSN Hardcoded', 'HIGH', 'Direct Impact'),
    ('6', 'Exported Apple Sign-In Callback', 'HIGH', 'Direct Impact'),
    ('7', 'Facebook Codeless Debug Logging', 'MEDIUM', 'Direct Impact'),
    ('8', 'GTM Preview Exported', 'MEDIUM', 'Direct Impact'),
    ('9', 'UXCam Session Recording', 'MEDIUM', 'Abuse Potential'),
    ('10', 'Task Hijacking (OneSignal)', 'MEDIUM', 'Abuse Potential'),
    ('11', 'TikTok SDK Data Collection', 'MEDIUM', 'Abuse Potential'),
    ('12', 'SYSTEM_ALERT_WINDOW Overlay', 'MEDIUM', 'Abuse Potential'),
    ('13', 'FCM Tokens in Plaintext', 'LOW-MOD', 'Abuse Potential'),
    ('14', 'HTTP Deep Link in Auto-Verify', 'LOW-MOD', 'Abuse Potential'),
    ('15', 'ClickUp API in Production', 'LOW', 'Informational'),
]

for i, (num, finding, sev, impact) in enumerate(findings):
    cells = table.rows[i+1].cells
    cells[0].text = num
    cells[1].text = finding
    cells[2].text = sev
    cells[3].text = impact

doc.add_paragraph('')
doc.add_paragraph('End of Report.')

doc.save('C:/Users/techmoiq/Desktop/Baity_Security_Findings.docx')
print('Done: Baity_Security_Findings.docx created')
